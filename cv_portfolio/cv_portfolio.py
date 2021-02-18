from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


document = Document()

# profile picture
document.add_picture(
    'me.jpg',
    width=Inches(2.0)
)

# profile details
name = input('What is your name?')
speak('Hello, ' + name + 'How are you today?')

speak('What is your phone number?')
phone_number = input('What is your phone number?')
email = input('What is your email')

document.add_paragraph(name + '|' + phone_number + "|" + email)

# about me
document.add_heading('About me')
document.add_paragraph(input('Tell me about yourself'))


# skills

def printSkill():
    skill = input('Enter a skill')
    p = document.add_paragraph(skill)
    p.style = 'List Bullet'


document.add_heading('Skills')
printSkill()

while True:
    has_more_skills = input('Do you have more skills? yes or no')
    if has_more_skills.lower() == 'yes':
        printSkill()
    else:
        break

document.save('cv.docx')
